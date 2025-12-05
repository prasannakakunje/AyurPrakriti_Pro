# AyurPrakriti_Pro_Mega.py
# FULL merged single-file Streamlit app
# - Admin system: YES
# - Full PDF fallback engine: YES
# - Large questionnaire set: YES
# - Full settings system: YES
#
# Save as AyurPrakriti_Pro_Mega.py and run:
# pip install streamlit reportlab matplotlib pandas python-docx passlib pyyaml pillow
# streamlit run AyurPrakriti_Pro_Mega.py
#
# NOTE: This file is long (split into 4 parts). Paste the 4 parts together in order.

import os, sys, json, shutil, logging, traceback
from pathlib import Path
from datetime import datetime, timedelta
from io import BytesIO
import yaml
import sqlite3
import base64

# UI / Data
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from passlib.context import CryptContext
from docx import Document
from PIL import Image

# ReportLab
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
    Table,
    TableStyle,
    PageBreak,
    KeepTogether,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

# ---------------- Directories & app metadata ----------------
APP_DIR = Path.home() / ".ayurprakriti_app"
APP_DIR.mkdir(parents=True, exist_ok=True)
FONTS_DIR = APP_DIR / "fonts"
FONTS_DIR.mkdir(parents=True, exist_ok=True)
TMP_DIR = APP_DIR / "tmp"
TMP_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR = APP_DIR / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = APP_DIR / "ayurprakriti.db"
CFG_PATH = APP_DIR / "config_rules.yaml"
LOG_PATH = APP_DIR / "app_debug.log"

# Try to copy logo from common container location if present
    try:
    except:
        pass

# Branding default (will be editable in UI)
BRAND = {
    "clinic_name": "Kakunje Wellness",
    "tagline": "Authentic Ayurveda | Modern Precision",
    "doctor": "Prof. Dr. Prasanna Kakunje, MD (Ayu), (PhD)",
    "address": "Janani Complex, Nagarakatte Road, Moodbidri, Karnataka",
    "phone": "+91-9483697676",
    "email": "prasanna@kakunje.com",
    "website": "https://kakunje.com",
    "accent_color": "#0F7A61",
}

# PDF watermark/footer defaults
WCONF = {
    "watermark_text": BRAND["clinic_name"],
    "watermark_opacity": 0.06,
    "show_footer_logo": True,
    "use_footer_signature": False,
    "page_number_format": "Page {page}",
    "footer_signature_file": str(APP_DIR / "signature.png"),
}

# ---------------- Logging ----------------
logger = logging.getLogger("ayurprakriti_mega")
if not logger.handlers:
    fh = logging.FileHandler(LOG_PATH)
    fh.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
    logger.addHandler(fh)
logger.setLevel(logging.INFO)

# ---------------- Config defaults (large questionnaire + mappings) ----------------
pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")

# Comprehensive question banks (expanded)
DEFAULT_CFG = {
    "meta": {
        "app_name": "AyurPrakriti Pro Mega",
        "version": "2.0",
        "author": "Generated",
    },
    "questions": {
        "prakriti": [
            # Expanded set (sample ~25). Real deployments can extend via YAML UI.
            {
                "id": "P1",
                "text": "Natural body frame: thin/slender",
                "weights": {"Vata": 1.0},
            },
            {
                "id": "P2",
                "text": "Tendency for dry, rough skin",
                "weights": {"Vata": 1.0},
            },
            {
                "id": "P3",
                "text": "Variable appetite / digestion",
                "weights": {"Vata": 1.0},
            },
            {
                "id": "P4",
                "text": "Light sleep, easily awakened",
                "weights": {"Vata": 1.0},
            },
            {
                "id": "P5",
                "text": "Quick, changeable mood",
                "weights": {"Vata": 0.9, "Pitta": 0.1},
            },
            {"id": "P6", "text": "Warm body/flush easily", "weights": {"Pitta": 1.0}},
            {
                "id": "P7",
                "text": "Strong appetite, tolerates spicy",
                "weights": {"Pitta": 1.0},
            },
            {
                "id": "P8",
                "text": "Ambitious, focused under pressure",
                "weights": {"Pitta": 0.8},
            },
            {"id": "P9", "text": "Calm, steady energy", "weights": {"Kapha": 1.0}},
            {
                "id": "P10",
                "text": "Good endurance and build",
                "weights": {"Kapha": 1.0},
            },
            {"id": "P11", "text": "Tendency to gain weight", "weights": {"Kapha": 1.0}},
            {
                "id": "P12",
                "text": "Slow digestion vs regular digestion",
                "weights": {"Kapha": 0.7, "Vata": 0.3},
            },
            {"id": "P13", "text": "Cold extremities often", "weights": {"Vata": 0.8}},
            {
                "id": "P14",
                "text": "Perspiration: sweats easily",
                "weights": {"Pitta": 0.7},
            },
            {
                "id": "P15",
                "text": "Memory: quick recall vs steady long-term",
                "weights": {"Vata": 0.5, "Kapha": 0.5},
            },
            {
                "id": "P16",
                "text": "Preference for warm foods",
                "weights": {"Vata": 0.6},
            },
            {
                "id": "P17",
                "text": "Tendency for oily skin",
                "weights": {"Pitta": 0.6, "Kapha": 0.4},
            },
            {
                "id": "P18",
                "text": "Joint stiffness when inactive",
                "weights": {"Kapha": 0.8},
            },
            {
                "id": "P19",
                "text": "Speech: fast vs slow",
                "weights": {"Vata": 0.7, "Kapha": 0.3},
            },
            {
                "id": "P20",
                "text": "Physical strength & stamina",
                "weights": {"Kapha": 0.7, "Pitta": 0.3},
            },
            {
                "id": "P21",
                "text": "Prone to allergies/congestion",
                "weights": {"Kapha": 0.7, "Pitta": 0.3},
            },
            {
                "id": "P22",
                "text": "Easily excited / enthusiastic",
                "weights": {"Vata": 0.7, "Pitta": 0.3},
            },
            {
                "id": "P23",
                "text": "Face color: reddish vs pale",
                "weights": {"Pitta": 0.8, "Kapha": 0.4},
            },
            {
                "id": "P24",
                "text": "Thirst level (high/low)",
                "weights": {"Pitta": 0.7, "Kapha": 0.3},
            },
            {
                "id": "P25",
                "text": "Tendency for constipation",
                "weights": {"Vata": 0.9},
            },
        ],
        "vikriti": [
            # Expanded vikriti sample (~20)
            {
                "id": "V1",
                "text": "Anxiety, restlessness today",
                "weights": {"Vata": 1.0},
            },
            {"id": "V2", "text": "Racing thoughts, insomnia", "weights": {"Vata": 1.0}},
            {"id": "V3", "text": "Cold hands/feet today", "weights": {"Vata": 0.8}},
            {
                "id": "V4",
                "text": "Excess heat, anger, irritability",
                "weights": {"Pitta": 1.0},
            },
            {
                "id": "V5",
                "text": "Acidity, heartburn, sour belching",
                "weights": {"Pitta": 1.0},
            },
            {
                "id": "V6",
                "text": "Red rashes or inflammation",
                "weights": {"Pitta": 1.0},
            },
            {
                "id": "V7",
                "text": "Heaviness, lethargy, sleepiness",
                "weights": {"Kapha": 1.0},
            },
            {
                "id": "V8",
                "text": "Congestion, phlegm, mucus",
                "weights": {"Kapha": 1.0},
            },
            {
                "id": "V9",
                "text": "Slow digestion, poor appetite",
                "weights": {"Kapha": 0.8},
            },
            {
                "id": "V10",
                "text": "Joint stiffness or swelling",
                "weights": {"Kapha": 0.7},
            },
            {
                "id": "V11",
                "text": "Excess thirst or dry mouth",
                "weights": {"Pitta": 0.6},
            },
            {
                "id": "V12",
                "text": "Loose stools or irregular digestion",
                "weights": {"Vata": 0.8},
            },
            {"id": "V13", "text": "Excess worrying today", "weights": {"Vata": 0.9}},
            {
                "id": "V14",
                "text": "Agitation or short temper",
                "weights": {"Pitta": 0.9},
            },
            {"id": "V15", "text": "Sleep fragmented", "weights": {"Vata": 0.8}},
            {
                "id": "V16",
                "text": "Sensation of heaviness in the head",
                "weights": {"Kapha": 0.7},
            },
            {"id": "V17", "text": "Excess sweating", "weights": {"Pitta": 0.5}},
            {"id": "V18", "text": "Reduced motivation", "weights": {"Kapha": 0.8}},
            {
                "id": "V19",
                "text": "Unusual cravings (salty/sweet)",
                "weights": {"Kapha": 0.6},
            },
            {
                "id": "V20",
                "text": "Irritable bowel symptoms",
                "weights": {"Pitta": 0.6, "Vata": 0.4},
            },
        ],
        "psychometric": [
            # Expanded personality-like items (10 pairs)
            {"id": "E1", "text": "Outgoing, enthusiastic"},
            {"id": "E6", "text": "Reserved, quiet"},
            {"id": "A1", "text": "Often critical"},
            {"id": "A6", "text": "Warm, sympathetic"},
            {"id": "C1", "text": "Organized, reliable"},
            {"id": "C6", "text": "Disorganized, careless"},
            {"id": "N1", "text": "Often anxious"},
            {"id": "N6", "text": "Emotionally stable"},
            {"id": "O1", "text": "Open to new ideas"},
            {"id": "O6", "text": "Conventional, prefers routine"},
        ],
    },
    "mappings": {
        "career_rules": {
            "Vata": ["Writer", "Designer", "Consultant - Creative", "Researcher"],
            "Pitta": ["Clinician", "Analyst", "Manager", "Engineer"],
            "Kapha": ["Teacher", "Counselor", "Hospitality", "HR", "Agriculture"],
        },
        "dosha_thresholds": {"mild": 55, "moderate": 70, "severe": 85},
    },
}

# If config file not present, write defaults
if not CFG_PATH.exists():
    with open(CFG_PATH, "w", encoding="utf-8") as f:
        yaml.safe_dump(DEFAULT_CFG, f, sort_keys=False)
# Load config
with open(CFG_PATH, "r", encoding="utf-8") as f:
    CONFIG = yaml.safe_load(f)

# ---------------- Database init ----------------
conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
cur = conn.cursor()
cur.executescript(
    """
CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT UNIQUE,
    display_name TEXT,
    password_hash TEXT,
    role TEXT DEFAULT 'clinician',
    created_at TEXT
);
CREATE TABLE IF NOT EXISTS patients (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT,
    age INTEGER,
    gender TEXT,
    contact TEXT,
    created_at TEXT
);
CREATE TABLE IF NOT EXISTS assessments (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    patient_id INTEGER,
    assessor TEXT,
    data_json TEXT,
    created_at TEXT,
    FOREIGN KEY(patient_id) REFERENCES patients(id)
);
"""
)
conn.commit()

# Create default admin if missing
cur.execute("SELECT COUNT(1) FROM users")
if cur.fetchone()[0] == 0:
    ph = pwd_context.hash("admin123")
    cur.execute(
        "INSERT INTO users (username, display_name, password_hash, role, created_at) VALUES (?,?,?,?,?)",
        ("admin", "Administrator", ph, "admin", datetime.now().isoformat()),
    )
    conn.commit()


# ---------------- Utility helpers ----------------
def verify_user(username, password):
    cur.execute(
        "SELECT password_hash, display_name, role FROM users WHERE username=?",
        (username,),
    )
    r = cur.fetchone()
    if not r:
        return False, None
    ph, display, role = r
    try:
        ok = pwd_context.verify(password, ph)
    except Exception:
        ok = False
    return ok, {"display_name": display, "role": role}


def create_patient(name, age, gender, contact):
    cur.execute(
        "INSERT INTO patients (name, age, gender, contact, created_at) VALUES (?,?,?,?,?)",
        (name, age, gender, contact, datetime.now().isoformat()),
    )
    conn.commit()
    return cur.lastrowid


def save_assessment(patient_id, assessor, data):
    cur.execute(
        "INSERT INTO assessments (patient_id, assessor, data_json, created_at) VALUES (?,?,?,?)",
        (
            patient_id,
            assessor,
            json.dumps(data, ensure_ascii=False),
            datetime.now().isoformat(),
        ),
    )
    conn.commit()
    return cur.lastrowid


def load_patients():
    return pd.read_sql_query("SELECT * FROM patients ORDER BY created_at DESC", conn)


def load_assessments(patient_id=None):
    if patient_id:
        return pd.read_sql_query(
            "SELECT * FROM assessments WHERE patient_id=? ORDER BY created_at DESC",
            conn,
            params=(patient_id,),
        )
    return pd.read_sql_query("SELECT * FROM assessments ORDER BY created_at DESC", conn)


# ---------------- Scoring functions (dosha, psych) ----------------
def score_dosha_from_answers(answers, question_list):
    totals = {"Vata": 0.0, "Pitta": 0.0, "Kapha": 0.0}
    for q in question_list:
        qid = q["id"]
        w = q.get("weights", {})
        val = answers.get(qid, 3)
        for d in totals:
            totals[d] += w.get(d, 0) * float(val)
    s = sum(totals.values())
    if s <= 0:
        return {k: round(100 / 3, 1) for k in totals}
    return {k: round((v / s) * 100, 1) for k, v in totals.items()}


def psychometric_tipiscale(answers):
    try:
        ext = (answers["E1"] + (8 - answers["E6"])) / 2.0
        agr = (((8 - answers["A1"]) + answers["A6"])) / 2.0
        con = (answers["C1"] + (8 - answers["C6"])) / 2.0
        emo = (answers["N1"] + (8 - answers["N6"])) / 2.0
        ope = (answers["O1"] + (8 - answers["O6"])) / 2.0
    except Exception:
        return {
            "Extraversion": 50,
            "Agreeableness": 50,
            "Conscientiousness": 50,
            "Emotionality": 50,
            "Openness": 50,
        }
    raw = {
        "Extraversion": ext,
        "Agreeableness": agr,
        "Conscientiousness": con,
        "Emotionality": emo,
        "Openness": ope,
    }
    return {k: round((v - 1) / 6 * 100, 1) for k, v in raw.items()}


# ---------------- Recommendation engines ----------------
def recommend_career(dosha_percent, psycho_pct, cfg=CONFIG):
    dom = max(dosha_percent, key=dosha_percent.get)
    base = cfg["mappings"]["career_rules"].get(dom, [])
    recs = []
    for r in base:
        score = 50
        if psycho_pct.get("Openness", 50) > 65 and (
            "Research" in r or "Creative" in r or "Writer" in r
        ):
            score += 10
        if psycho_pct.get("Conscientiousness", 50) > 65 and (
            "Manager" in r or "Engineer" in r
        ):
            score += 8
        if psycho_pct.get("Extraversion", 50) > 60 and (
            "Clinician" in r or "Teacher" in r
        ):
            score += 6
        recs.append(
            {
                "role": r,
                "score": score,
                "reason": f"Matches dominant {dom} + personality cues.",
            }
        )
    if psycho_pct.get("Openness", 50) > 70 and not any(
        "Research" in x["role"] for x in recs
    ):
        recs.append(
            {
                "role": "Research & Innovation",
                "score": 65,
                "reason": "High openness suggests research fit.",
            }
        )
    return sorted(recs, key=lambda x: -x["score"])


def recommend_relationship(dosha_pct, psycho_pct):
    tips = []
    dom = max(dosha_pct, key=dosha_pct.get)
    if dom == "Vata":
        tips.append(
            (
                "Stability & routines",
                "Vata benefits from grounding, predictable routines; short daily check-ins help.",
            )
        )
    if dom == "Pitta":
        tips.append(
            (
                "Cooling communication",
                "Pause before responding and use neutral language during disagreements.",
            )
        )
    if dom == "Kapha":
        tips.append(
            (
                "Introduce small novelty",
                "Gentle new activities reduce inertia and boost engagement.",
            )
        )
    if psycho_pct.get("Agreeableness", 50) < 40:
        tips.append(
            (
                "Reflective listening",
                "Summarize what partner said before giving your view.",
            )
        )
    if psycho_pct.get("Emotionality", 50) > 60:
        tips.append(
            (
                "Emotion regulation",
                "Use 3-minute breathing or journaling before difficult talks.",
            )
        )
    return tips


def recommend_health(dosha_pct, vikriti_pct, cfg=CONFIG):
    dom = max(dosha_pct, key=dosha_pct.get)
    rec = {"diet": [], "lifestyle": [], "herbs": [], "severity": {}}
    for d in dosha_pct:
        score = round((dosha_pct[d] + vikriti_pct.get(d, 0)) / 2, 1)
        if score >= cfg["mappings"]["dosha_thresholds"]["severe"]:
            rec["severity"][d] = "severe"
        elif score >= cfg["mappings"]["dosha_thresholds"]["moderate"]:
            rec["severity"][d] = "moderate"
        elif score >= cfg["mappings"]["dosha_thresholds"]["mild"]:
            rec["severity"][d] = "mild"
        else:
            rec["severity"][d] = "balanced"
    if dom == "Vata":
        rec["diet"] = [
            "Warm, cooked meals; include healthy oils; regular meal timings; avoid iced drinks first thing."
        ]
        rec["lifestyle"] = [
            "Daily warm oil massage (Abhyanga) 5–10 min; grounding morning routine; consistent sleep schedule."
        ]
        rec["herbs"] = ["Ashwagandha (under clinician guidance), Bala for strength."]
    if dom == "Pitta":
        rec["diet"] = [
            "Cooling foods; reduce spicy, fried and fermented foods; include bitter greens."
        ]
        rec["lifestyle"] = [
            "Avoid midday heat; cooling pranayama; calm, regular breaks."
        ]
        rec["herbs"] = ["Amla, Guduchi (clinician review)."]
    if dom == "Kapha":
        rec["diet"] = ["Light, warm, slightly astringent foods; reduce dairy & sweets."]
        rec["lifestyle"] = [
            "Stimulating exercise 30–60 min daily; vary routine; dry massage (udvartana)."
        ]
        rec["herbs"] = ["Trikatu, Guggulu (clinician supervision)."]
    return rec


# ---------------- Charting helpers (bars + radar) ----------------
def _make_bar_chart(series: dict, title: str, filename: Path):
    plt.close("all")
    keys = list(series.keys())
    vals = [series[k] for k in keys]
    fig, ax = plt.subplots(figsize=(6, 2.6))
    palette = ["#6fbf73", "#f5a623", "#6fb0d9"]
    bars = ax.bar(keys, vals, color=palette[: len(keys)])
    ax.set_ylim(0, 100)
    ax.set_ylabel("Percent")
    ax.set_title(title, fontsize=10)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            v + 1,
            f"{v}%",
            ha="center",
            va="bottom",
            fontsize=8,
        )
    fig.tight_layout()
    fig.savefig(filename, dpi=150)
    plt.close(fig)


def make_radar_chart(prakriti, vikriti, filename: Path, title="Prakriti vs Vikriti"):
    labels = list(prakriti.keys())
    n = len(labels)
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    vals1 = [prakriti[l] for l in labels]
    vals2 = [vikriti.get(l, 0) for l in labels]
    vals1 += vals1[:1]
    vals2 += vals2[:1]
    angles += angles[:1]
    fig = plt.figure(figsize=(4.2, 4.2))
    ax = fig.add_subplot(111, polar=True)
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.plot(angles, vals1, linewidth=2, label="Prakriti")
    ax.fill(angles, vals1, alpha=0.25)
    ax.plot(angles, vals2, linewidth=2, label="Vikriti")
    ax.fill(angles, vals2, alpha=0.12)
    ax.set_thetagrids(np.degrees(angles[:-1]), labels)
    ax.set_ylim(0, 100)
    ax.set_title(title, pad=10)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1))
    plt.tight_layout()
    fig.savefig(filename, dpi=150)
    plt.close(fig)


# ---------------- Fonts registration (DejaVu) ----------------
DEJAVU_PATH = None
_fonts = list(FONTS_DIR.glob("DejaVuSans*.ttf"))
if _fonts:
    DEJAVU_PATH = str(_fonts[0])
else:
    for cand in [
        r"C:\Windows\Fonts\DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/DejaVuSans.ttf",
    ]:
        if os.path.exists(cand):
            DEJAVU_PATH = cand
            break
if DEJAVU_PATH:
    try:
        pdfmetrics.registerFont(TTFont("DejaVuSans", DEJAVU_PATH))
        logger.info("Registered DejaVu font: %s", DEJAVU_PATH)
    except Exception:
        logger.exception("Failed to register DejaVu")


# ---------------- Plain-language "WOW" advice generator ----------------
def generate_wow_advice(
    patient, prakriti_pct, vikriti_pct, psych_pct, career_recs, rel_tips, health_recs
):
    dom = max(prakriti_pct, key=prakriti_pct.get)
    current = max(vikriti_pct, key=vikriti_pct.get)
    hero = (
        f"{patient.get('name','You')} — you have {dom}-style strengths: creativity, quick thinking, and unique energy. "
        f"Right now you may feel {('slow & heavy' if current=='Kapha' else 'scattered & anxious' if current=='Vata' else 'hot & impatient')}."
    )
    plan_lines = [
        "90-day transformation plan (small actions -> identity change):",
        "Day 1: Identity pledge — write one line: 'I am someone who finishes what they start with calm focus.'",
        "Days 1–21: Core daily ritual — warm water, 5–10 min oil rub or warm stretch, two focused 60–90 min blocks.",
        "Weeks 4–12: Publish one small project every 2–3 weeks and get feedback.",
        "Accountability: pick a peer for weekly 2-min check-ins over 12 weeks.",
        "Measure: morning energy 1–5 and sleep time daily; review at day 14, 45, 90.",
    ]
    plan = "\n".join(plan_lines)
    habit_stack = "\n".join(
        [
            "Life-changing habit stack (15–25 min total):",
            "A) Warm water + 2 min breathing (inhale 4s / exhale 6s).",
            "B) 5–10 min oil massage or 10 min stretching.",
            "C) One 60–90 min focused work block (timer on).",
            "D) Evening reflection: list 2 wins and 1 tomorrow task.",
        ]
    )
    wow_tips = "\n".join(
        [
            "- Reduce decision fatigue: limit morning choices to 3 (clothes/breakfast).",
            "- Ship every week: a tiny deliverable that builds momentum.",
            "- Use '2-minute accountability' with a friend; micro-commitments scale.",
            "- Reassess and tweak after 14 days — small changes compound.",
        ]
    )
    checklist = "\n".join(
        [
            "ONE-PAGE ACTION CHECKLIST",
            "- Morning: warm water + 2 min breathing + 5–10 min oil rub/stretch",
            "- Work: 2 focused blocks (60–90 min each). Timer ON.",
            "- Movement: 25–35 min daily walk / yoga.",
            "- Evening: light dinner by 8 pm; reflect on 2 wins.",
            "- Weekly: share a small project and plan next week (20 min).",
            "- Accountability: weekly check-in with chosen peer for 12 weeks.",
        ]
    )
    doctor_note = (
        "Doctor's note: Begin the 'Start today' actions now. Small consistent changes matter more than rare big efforts. "
        "We will review progress at 2 weeks and refine the plan."
    )
    return {
        "hero": hero,
        "plan": plan,
        "habit_stack": habit_stack,
        "wow_tips": wow_tips,
        "checklist": checklist,
        "doctor_note": doctor_note,
    }


# ---------------- One-page Action Plan PDF ----------------
def onepage_actionplan_pdf(patient, checklist_text, hero_text):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    left = 20 * mm
    y = A4[1] - 30 * mm
    try:
        if DEJAVU_PATH:
            c.setFont("DejaVuSans", 14)
        else:
            c.setFont("Helvetica-Bold", 14)
    except:
        c.setFont("Helvetica-Bold", 14)
    c.drawString(left, y, BRAND["clinic_name"])
    y -= 8 * mm
    c.setFont("Helvetica", 10)
    c.drawString(left, y, hero_text)
    y -= 9 * mm
    c.setFont("Helvetica", 10)
    for line in checklist_text.split("\n"):
        if not line.strip():
            continue
        if line.startswith("- "):
            c.drawString(left + 4 * mm, y, "\u2022 " + line[2:])
        else:
            c.drawString(left, y, line)
        y -= 7 * mm
        if y < 30 * mm:
            c.showPage()
            y = A4[1] - 30 * mm
    c.setFont("Helvetica", 8)
    c.drawString(
        left, 12 * mm, f"{BRAND['clinic_name']} — {BRAND['phone']} — {BRAND['email']}"
    )
    c.save()
    buf.seek(0)
    return buf


# ---------------- Simple text wrapper ----------------
def _wrap_text_simple(text, chars_per_line=95):
    words = str(text).split()
    lines = []
    cur = ""
    for w in words:
        if len(cur) + len(w) + 1 <= chars_per_line:
            cur = cur + (" " if cur else "") + w
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


# ---------------- PDF builders: platypus branded + fallback canvas ----------------
def branded_pdf_report(
    patient,
    prakriti_pct,
    vikriti_pct,
    psych_pct,
    career_recs,
    rel_tips,
    health_recs,
    include_appendix=False,
    report_id=None,
    wconf=None,
    wow=None,
):
    if wconf is None:
        wconf = WCONF
    # generate charts
    p1 = TMP_DIR / f"prakriti_{int(datetime.now().timestamp())}.png"
    p2 = TMP_DIR / f"vikriti_{int(datetime.now().timestamp())}.png"
    p3 = TMP_DIR / f"psych_{int(datetime.now().timestamp())}.png"
    radar = TMP_DIR / f"radar_{int(datetime.now().timestamp())}.png"
    try:
        _make_bar_chart(prakriti_pct, "Prakriti (constitutional %)", p1)
        _make_bar_chart(vikriti_pct, "Vikriti (today %)", p2)
        _make_bar_chart(psych_pct, "Psychometric (approx %)", p3)
        make_radar_chart(prakriti_pct, vikriti_pct, radar)
    except Exception:
        logger.exception("Chart generation failed")

    try:
        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
        )
        styles = getSampleStyleSheet()
        base_font = "DejaVuSans" if DEJAVU_PATH else "Helvetica"
        accent = colors.HexColor(BRAND["accent_color"])
        styles.add(
            ParagraphStyle(
                name="AP_Title",
                fontName=base_font,
                fontSize=18,
                leading=22,
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(name="AP_Small", fontName=base_font, fontSize=9, leading=11)
        )
        styles.add(
            ParagraphStyle(
                name="AP_Heading",
                fontName=base_font,
                fontSize=12,
                leading=14,
                spaceBefore=8,
                spaceAfter=4,
                textColor=accent,
            )
        )
        styles.add(
            ParagraphStyle(name="AP_Body", fontName=base_font, fontSize=10, leading=13)
        )
        styles.add(
            ParagraphStyle(
                name="AP_Bullet",
                fontName=base_font,
                fontSize=10,
                leading=12,
                leftIndent=12,
                bulletIndent=6,
            )
        )

        flow = []
        # Cover / Hero
        flow.append(Spacer(1, 6))
        # header with logo if exists
        logo_path = APP_DIR / "logo.png"
        if logo_path.exists():
            try:
                img = RLImage(str(logo_path), width=40 * mm, height=40 * mm)
                clinic_info = Paragraph(
                    f"<b>{BRAND['clinic_name']}</b><br/>{BRAND['tagline']}<br/><font size=9>{BRAND['website']}</font>",
                    styles["AP_Body"],
                )
                header_t = Table([[img, clinic_info]], colWidths=[45 * mm, 120 * mm])
                header_t.setStyle(
                    TableStyle(
                        [
                            ("VALIGN", (0, 0), (1, 0), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ]
                    )
                )
                flow.append(header_t)
            except Exception:
                flow.append(
                    Paragraph(
                        f"<b>{BRAND['clinic_name']}</b><br/>{BRAND['tagline']}",
                        styles["AP_Title"],
                    )
                )
        else:
            flow.append(
                Paragraph(
                    f"<b>{BRAND['clinic_name']}</b><br/>{BRAND['tagline']}",
                    styles["AP_Title"],
                )
            )
        flow.append(Spacer(1, 6))
        # hero name + short insight
        flow.append(
            Paragraph(
                f"<b>{patient.get('name','Patient Name')}</b>", styles["AP_Title"]
            )
        )
        if wow and wow.get("hero"):
            flow.append(Paragraph(wow.get("hero"), styles["AP_Body"]))
        flow.append(Spacer(1, 8))

        # badges row
        badges = [
            Paragraph(
                f"<b>Dominant</b><br/>{max(prakriti_pct, key=prakriti_pct.get)}",
                styles["AP_Body"],
            ),
            Paragraph(
                f"<b>Current</b><br/>{max(vikriti_pct, key=vikriti_pct.get)}",
                styles["AP_Body"],
            ),
            Paragraph(
                f"<b>Top career</b><br/>{career_recs[0]['role'] if career_recs else '-'}",
                styles["AP_Body"],
            ),
        ]
        t_badges = Table(
            [[badges[0], badges[1], badges[2]]], colWidths=[60 * mm, 60 * mm, 60 * mm]
        )
        t_badges.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        flow.append(t_badges)
        flow.append(Spacer(1, 8))

        # radar on cover (if generated)
        if radar.exists():
            try:
                rimg = RLImage(str(radar), width=120 * mm, height=120 * mm)
                flow.append(rimg)
                flow.append(Spacer(1, 8))
            except Exception:
                pass

        # doctor note + signature
        if wow and wow.get("doctor_note"):
            flow.append(
                Paragraph(f"<i>{wow.get('doctor_note')}</i>", styles["AP_Body"])
            )
            sig = APP_DIR / "signature.png"
            if sig.exists():
                try:
                    s_img = RLImage(str(sig), width=40 * mm, height=15 * mm)
                    flow.append(Spacer(1, 4))
                    flow.append(s_img)
                except:
                    pass
        flow.append(PageBreak())

        # Executive summary & charts
        flow.append(Paragraph("Executive summary", styles["AP_Heading"]))
        flow.append(
            Paragraph(
                "This report summarises constitutional profile (Prakriti), current imbalances (Vikriti), psychometric snapshot and prioritized recommendations.",
                styles["AP_Body"],
            )
        )
        flow.append(Spacer(1, 8))

        # Add bar charts (if created)
        try:
            if p1.exists() and p2.exists():
                img1 = RLImage(str(p1), width=85 * mm, height=45 * mm)
                img2 = RLImage(str(p2), width=85 * mm, height=45 * mm)
                flow.append(Table([[img1, img2]], colWidths=[90 * mm, 90 * mm]))
                flow.append(Spacer(1, 6))
            if p3.exists():
                img3 = RLImage(str(p3), width=160 * mm, height=35 * mm)
                flow.append(img3)
                flow.append(Spacer(1, 6))
        except Exception:
            logger.exception("Adding chart images failed")

        # Prakriti/Vikriti Tables
        flow.append(
            Paragraph("Prakriti — percentage distribution", styles["AP_Heading"])
        )
        pp = [[k, f"{v} %"] for k, v in prakriti_pct.items()]
        tpp = Table(pp, colWidths=[80 * mm, 80 * mm])
        tpp.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        flow.append(tpp)
        flow.append(Spacer(1, 6))
        flow.append(
            Paragraph("Vikriti — percentage distribution (today)", styles["AP_Heading"])
        )
        vp = [[k, f"{v} %"] for k, v in vikriti_pct.items()]
        tvp = Table(vp, colWidths=[80 * mm, 80 * mm])
        tvp.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        flow.append(tvp)
        flow.append(Spacer(1, 8))

        # Priority action strip (Start today / This week / This month)
        priority = [
            (
                "Start today",
                "Warm water on waking; 5–10 min warm oil rub or stretch; one focused 60–90 min block",
            ),
            (
                "This week",
                "Add second focused block; daily 20–35 min walk; start micro-project",
            ),
            (
                "This month",
                "Finish + share one creative project; weekly accountability",
            ),
        ]
        cols_cells = []
        for title, text in priority:
            cols_cells.append(
                Paragraph(f"<b>{title}</b><br/>{text}", styles["AP_Body"])
            )
        strip_tbl = Table([cols_cells], colWidths=[60 * mm, 60 * mm, 60 * mm])
        strip_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.Color(0.95, 0.98, 0.95)),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.lightgrey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        flow.append(strip_tbl)
        flow.append(Spacer(1, 8))

        # Recommendations short blocks
        flow.append(Paragraph("Recommendations — prioritized", styles["AP_Heading"]))
        flow.append(Paragraph("<b>Career</b>:", styles["AP_Body"]))
        for cr in career_recs[:6]:
            flow.append(
                Paragraph(
                    f"• <b>{cr['role']}</b> — score {cr['score']}. <i>{cr['reason']}</i>",
                    styles["AP_Bullet"],
                )
            )
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("<b>Relationship tips</b>:", styles["AP_Body"]))
        for t in rel_tips:
            flow.append(Paragraph(f"• <b>{t[0]}</b> — {t[1]}", styles["AP_Bullet"]))
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("<b>Health (diet & lifestyle)</b>:", styles["AP_Body"]))
        for d in health_recs.get("diet", []):
            flow.append(Paragraph(f"• {d}", styles["AP_Bullet"]))
        for l in health_recs.get("lifestyle", []):
            flow.append(Paragraph(f"• {l}", styles["AP_Bullet"]))
        flow.append(Spacer(1, 8))

        # Appendices / wow plan
        if include_appendix and wow:
            flow.append(PageBreak())
            flow.append(
                Paragraph(
                    "APPENDIX — Transformation & Practical Plan", styles["AP_Heading"]
                )
            )
            flow.append(Spacer(1, 6))
            flow.append(
                Paragraph("<b>90-day transformation plan</b>", styles["AP_Body"])
            )
            for line in wow["plan"].split("\n"):
                flow.append(Paragraph(line, styles["AP_Body"]))
            flow.append(Spacer(1, 6))
            flow.append(Paragraph("<b>Daily habit stack</b>", styles["AP_Body"]))
            for line in wow["habit_stack"].split("\n"):
                flow.append(Paragraph(line, styles["AP_Body"]))
            flow.append(Spacer(1, 6))
            flow.append(Paragraph("<b>Concrete tips</b>", styles["AP_Body"]))
            for line in wow["wow_tips"].split("\n"):
                flow.append(Paragraph(line, styles["AP_Body"]))
            flow.append(Spacer(1, 6))
            flow.append(Paragraph("<b>One-page checklist</b>", styles["AP_Body"]))
            for line in wow["checklist"].split("\n"):
                flow.append(Paragraph(line, styles["AP_Body"]))

        # Contact/footer
        flow.append(Spacer(1, 12))
        contact_par = (
            f"{BRAND.get('clinic_name')} — {BRAND.get('doctor')} — {BRAND.get('phone')}"
        )
        flow.append(Paragraph(contact_par, styles["AP_Small"]))
        flow.append(Paragraph(BRAND.get("address", ""), styles["AP_Small"]))

        # Watermark and footer function
        def _draw_page_footer_and_watermark(canvas_obj, doc_obj):
            try:
                canvas_obj.saveState()
                W, H = A4
                try:
                    if DEJAVU_PATH:
                        canvas_obj.setFont("DejaVuSans", 36)
                    else:
                        canvas_obj.setFont("Helvetica-Bold", 36)
                except Exception:
                    canvas_obj.setFont("Helvetica-Bold", 36)
                opacity = float(wconf.get("watermark_opacity", 0.06))
                try:
                    canvas_obj.setFillAlpha(opacity)
                except Exception:
                    canvas_obj.setFillColorRGB(0.7, 0.7, 0.7)
                canvas_obj.translate(W / 2.0, H / 2.0)
                canvas_obj.rotate(30)
                canvas_obj.drawCentredString(
                    0, 0, wconf.get("watermark_text", BRAND["clinic_name"])
                )
                canvas_obj.restoreState()
            except Exception:
                logger.exception("Watermark draw failed")
            try:
                canvas_obj.saveState()
                footer_y = 18 * mm
                canvas_obj.setStrokeColor(colors.lightgrey)
                canvas_obj.setLineWidth(0.5)
                canvas_obj.line(18 * mm, footer_y + 8, (A4[0] - 18 * mm), footer_y + 8)
                logo_path_local = APP_DIR / "logo.png"
                signature_path = (
                    Path(wconf.get("footer_signature_file", ""))
                    if wconf.get("footer_signature_file")
                    else None
                )
                x = 20 * mm
                if wconf.get("show_footer_logo", True) and logo_path_local.exists():
                    try:
                        reader = ImageReader(str(logo_path_local))
                        iw, ih = reader.getSize()
                        target_h = 10 * mm
                        scale = target_h / ih
                        canvas_obj.drawImage(
                            str(logo_path_local),
                            x,
                            footer_y - 2,
                            width=iw * scale,
                            height=ih * scale,
                            mask="auto",
                        )
                        x += (iw * scale) + 4
                    except Exception:
                        logger.exception("Footer logo draw error")
                elif (
                    wconf.get("use_footer_signature", False)
                    and signature_path
                    and signature_path.exists()
                ):
                    try:
                        reader = ImageReader(str(signature_path))
                        iw, ih = reader.getSize()
                        target_h = 10 * mm
                        scale = target_h / ih
                        canvas_obj.drawImage(
                            str(signature_path),
                            x,
                            footer_y - 2,
                            width=iw * scale,
                            height=ih * scale,
                            mask="auto",
                        )
                        x += (iw * scale) + 4
                    except Exception:
                        logger.exception("Footer signature draw error")
                try:
                    if DEJAVU_PATH:
                        canvas_obj.setFont("DejaVuSans", 8)
                    else:
                        canvas_obj.setFont("Helvetica", 8)
                except Exception:
                    canvas_obj.setFont("Helvetica", 8)
                contact_line = f"{BRAND.get('clinic_name')} — {BRAND.get('doctor')} — {BRAND.get('phone')} — {BRAND.get('email')}"
                canvas_obj.setFillColor(colors.HexColor("#444444"))
                canvas_obj.drawString(
                    18 * mm if x < 18 * mm + 2 else x, footer_y, contact_line
                )
                fmt = wconf.get("page_number_format", "Page {page}")
                try:
                    page_num = canvas_obj.getPageNumber()
                except Exception:
                    page_num = doc_obj.page
                if "{total}" in fmt:
                    page_text = fmt.replace("{page}", "%d").replace("{total}", "%d") % (
                        page_num,
                        page_num,
                    )
                else:
                    page_text = fmt.format(page=page_num)
                canvas_obj.drawRightString(A4[0] - 18 * mm, footer_y, page_text)
                canvas_obj.restoreState()
            except Exception:
                logger.exception("Footer drawing failed")

        doc.build(
            flow,
            onFirstPage=_draw_page_footer_and_watermark,
            onLaterPages=_draw_page_footer_and_watermark,
        )
        buf.seek(0)
        # cleanup temp images
        for p in [p1, p2, p3, radar]:
            try:
                if p.exists():
                    p.unlink()
            except:
                pass
        return buf
    except Exception:
        tb = traceback.format_exc()
        logger.exception("Platypus build failed: %s", tb)
        snippet = tb[:1200]
        return _fallback_canvas_pdf(
            patient,
            prakriti_pct,
            vikriti_pct,
            psych_pct,
            career_recs,
            rel_tips,
            health_recs,
            error_text=snippet,
            include_appendix=include_appendix,
            report_id=report_id,
            wconf=wconf,
            wow=wow,
        )


def _fallback_canvas_pdf(
    patient,
    prakriti_pct,
    vikriti_pct,
    psych_pct,
    career_recs,
    rel_tips,
    health_recs,
    error_text=None,
    include_appendix=False,
    report_id=None,
    wconf=None,
    wow=None,
):
    if wconf is None:
        wconf = WCONF
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    left = 18 * mm
    top = A4[1] - 18 * mm
    y = top
    try:
        logo_path = APP_DIR / "logo.png"
        if logo_path.exists():
            try:
                reader = ImageReader(str(logo_path))
                iw, ih = reader.getSize()
                scale = min((36 * mm) / iw, (36 * mm) / ih, 1.0)
                c.drawImage(
                    str(logo_path),
                    left,
                    y - (36 * mm),
                    width=iw * scale,
                    height=ih * scale,
                    mask="auto",
                )
            except Exception:
                pass
        c.setFont("Helvetica-Bold", 13)
        c.drawString(left, y - 6, BRAND.get("clinic_name", ""))
        y -= 24
        c.setFont("Helvetica", 9)
        c.drawString(left, y, f"Patient: {patient.get('name','')}")
        c.drawString(left + 220, y, f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        y -= 16
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left, y, "Prakriti:")
        y -= 12
        c.setFont("Helvetica", 9)
        for k, v in prakriti_pct.items():
            c.drawString(left + 6, y, f"{k}: {v} %")
            y -= 10
            if y < 60 * mm:
                c.showPage()
                y = top
        c.drawString(left, y, "Vikriti:")
        y -= 12
        for k, v in vikriti_pct.items():
            c.drawString(left + 6, y, f"{k}: {v} %")
            y -= 10
            if y < 60 * mm:
                c.showPage()
                y = top
        c.drawString(left, y, "Top recommendations:")
        y -= 12
        for cr in career_recs[:10]:
            c.drawString(left + 6, y, f"- {cr.get('role')} (score {cr.get('score')})")
            y -= 10
            if y < 60 * mm:
                c.showPage()
                y = top
        if include_appendix and wow:
            c.showPage()
            y = top
            c.setFont("Helvetica-Bold", 12)
            c.drawString(left, y, "APPENDIX — Transformation Plan")
            y -= 14
            c.setFont("Helvetica", 9)
            for line in wow.get("plan", "").split("\n"):
                c.drawString(left, y, line)
                y -= 10
                if y < 40 * mm:
                    c.showPage()
                    y = top
            for line in wow.get("habit_stack", "").split("\n"):
                c.drawString(left, y, line)
                y -= 10
                if y < 40 * mm:
                    c.showPage()
                    y = top
        if error_text:
            c.showPage()
            y = top
            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Report engine error (short):")
            y -= 14
            c.setFont("Helvetica", 8)
            for line in _wrap_text_simple(error_text, 120):
                c.drawString(left, y, line)
                y -= 8
                if y < 30 * mm:
                    c.showPage()
                    y = top
        # footer
        footer_y = 18 * mm
        c.setStrokeColor(colors.lightgrey)
        c.line(18 * mm, footer_y + 8, (A4[0] - 18 * mm), footer_y + 8)
        if wconf.get("show_footer_logo", True) and (APP_DIR / "logo.png").exists():
            try:
                reader = ImageReader(str(APP_DIR / "logo.png"))
                iw, ih = reader.getSize()
                target_h = 10 * mm
                scale = target_h / ih
                c.drawImage(
                    str(APP_DIR / "logo.png"),
                    left,
                    footer_y - 2,
                    width=iw * scale,
                    height=ih * scale,
                    mask="auto",
                )
            except:
                pass
        c.setFont("Helvetica", 8)
        c.drawString(
            left + 40,
            footer_y,
            f"{BRAND['clinic_name']} — {BRAND['phone']}",
        )
        c.save()
        buf.seek(0)
        return buf
    except Exception as e:
        # Log the exception and return a minimal error PDF so the caller still receives a BytesIO object.
        logger.exception("Fallback PDF generation failed: %s", e)
        err_buf = BytesIO()
        ec = canvas.Canvas(err_buf, pagesize=A4)
        ec.setFont("Helvetica-Bold", 12)
        ec.drawString(18 * mm, A4[1] - 30 * mm, "Error generating report")
        ec.setFont("Helvetica", 9)
        msg = str(e) if e else "Unknown error"
        for i, line in enumerate(_wrap_text_simple(msg, 90)):
            ec.drawString(18 * mm, A4[1] - (40 * mm + i * 8), line)
        if error_text:
            # include the shorter engine error if provided
            for j, line in enumerate(_wrap_text_simple(error_text, 90)):
                ec.drawString(18 * mm, A4[1] - (60 * mm + (i + j + 1) * 8), line)
        ec.save()
        err_buf.seek(0)
        return err_buf


def make_ics_followup(patient_name, days=7):
    start = (datetime.now() + timedelta(days=days)).strftime("%Y%m%dT090000")
    dtstamp = datetime.now().strftime("%Y%m%dT%H%M00")
    ics = f"BEGIN:VCALENDAR\nVERSION:2.0\nBEGIN:VEVENT\nDTSTAMP:{dtstamp}\nDTSTART:{start}\nSUMMARY:Follow-up — {patient_name}\nDESCRIPTION:Review Ayurveda plan and progress.\nEND:VEVENT\nEND:VCALENDAR"
    return ics.encode("utf-8")


# ---------------- Streamlit UI start ----------------
st.set_page_config(page_title=CONFIG["meta"]["app_name"], layout="wide")
st.markdown(
    "<style>section[data-testid='stSidebar'] {background-color: #f7f7fa}</style>",
    unsafe_allow_html=True,
)
header_col, right_col = st.columns([6, 4])
with header_col:
    st.title(CONFIG["meta"]["app_name"])
    st.caption(
        f"Version {CONFIG['meta'].get('version','')} — {CONFIG['meta'].get('author','')}"
    )
with right_col:
    st.write(datetime.now().strftime("%Y-%m-%d %H:%M"))

# ---------------- Authentication sidebar ----------------
st.sidebar.subheader("Login")
username = st.sidebar.text_input("Username")
password = st.sidebar.text_input("Password", type="password")
if "auth" not in st.session_state:
    st.session_state.auth = False
if st.sidebar.button("Login"):
    ok, info = verify_user(username, password)
    if ok:
        st.session_state.auth = True
        st.session_state.user = username
        st.session_state.user_info = info
        st.sidebar.success(f"Welcome {info['display_name']} ({info['role']})")
    else:
        st.sidebar.error("Invalid username/password")
if not st.session_state.auth:
    st.info("Please login from the sidebar (default admin / admin123).")
    st.stop()
st.sidebar.markdown("---")
st.sidebar.write("Role: " + st.session_state.user_info.get("role", "clinician"))

# ---------------- App tabs ----------------
tabs = st.tabs(
    ["Patient Registry", "New Assessment", "Clinician Dashboard", "Config & Export"]
)

# ----- Tab 1: Patient Registry -----
with tabs[0]:
    st.header("Patient Registry")
    with st.expander("Create new patient"):
        with st.form("new_patient_form"):
            pname = st.text_input("Full name")
            page = st.number_input("Age", min_value=0, max_value=120, value=30)
            pgender = st.selectbox(
                "Gender", ["Male", "Female", "Other", "Prefer not to say"]
            )
            pcontact = st.text_input("Contact (phone/email)")
            if st.form_submit_button("Create patient"):
                if not pname:
                    st.warning("Name required")
                else:
                    pid = create_patient(pname, page, pgender, pcontact)
                    st.success(f"Patient created (id: {pid})")
    st.write("### All patients")
    patients_df = load_patients()
    st.dataframe(patients_df)

# ----- Tab 2: New Assessment -----
with tabs[1]:
    st.header("New Assessment — Prakriti, Vikriti & Psychometrics")
    patients = load_patients()
    if patients.empty:
        st.info("No patients yet. Create one in Patient Registry.")
        st.stop()
    psel = st.selectbox(
        "Select patient",
        options=patients["id"].tolist(),
        format_func=lambda x: f"{int(x)} - {patients.loc[patients['id']==x,'name'].values[0]}",
    )
    patient_row = patients[patients["id"] == psel].iloc[0].to_dict()
    st.markdown(
        f"**Patient:** {patient_row['name']} | Age: {patient_row['age']} | Gender: {patient_row['gender']}"
    )
    st.markdown("---")

    # Render Prakriti questions dynamically (two-column)
    pr_qs = CONFIG["questions"]["prakriti"]
    pr_answers = {}
    cols = st.columns(2)
    for i, q in enumerate(pr_qs):
        with cols[i % 2]:
            pr_answers[q["id"]] = st.slider(q["text"], 1, 5, 3, key=f"pr_{q['id']}")
    # Vikriti (three-column)
    vk_qs = CONFIG["questions"]["vikriti"]
    vk_answers = {}
    cols = st.columns(3)
    for i, q in enumerate(vk_qs):
        with cols[i % 3]:
            vk_answers[q["id"]] = st.slider(q["text"], 1, 5, 1, key=f"vk_{q['id']}")
    # Psychometric (two-column, 1-7)
    psy_qs = CONFIG["questions"]["psychometric"]
    psy_answers = {}
    cols = st.columns(2)
    for i, q in enumerate(psy_qs):
        with cols[i % 2]:
            psy_answers[q["id"]] = st.slider(q["text"], 1, 7, 4, key=f"psy_{q['id']}")
    st.markdown("---")

    show_long_preview = st.checkbox(
        "Show long recommendations on screen (preview)", value=False
    )
    if st.button("Compute & Save Assessment"):
        # compute scores
        prak_pct = score_dosha_from_answers(pr_answers, pr_qs)
        vik_pct = score_dosha_from_answers(vk_answers, vk_qs)
        psych_pct = psychometric_tipiscale(psy_answers)
        career = recommend_career(prak_pct, psych_pct)
        rel = recommend_relationship(prak_pct, psych_pct)
        health = recommend_health(prak_pct, vik_pct)
        payload = {
            "patient": patient_row,
            "prakriti_answers": pr_answers,
            "vikriti_answers": vk_answers,
            "psych_answers": psy_answers,
            "prakriti_pct": prak_pct,
            "vikriti_pct": vik_pct,
            "psych_pct": psych_pct,
            "career_recs": career,
            "relationship_tips": rel,
            "health_recs": health,
            "created_at": datetime.now().isoformat(),
        }
        aid = save_assessment(patient_row["id"], st.session_state.user, payload)
        # generate wow advice and attach
        wow = generate_wow_advice(
            patient_row, prak_pct, vik_pct, psych_pct, career, rel, health
        )
        payload["wow"] = wow
        st.session_state["last_assessment"] = payload
        st.session_state["last_aid"] = aid
        st.success(f"Assessment saved (id: {aid})")

    # if assessment exists in session, show preview & downloads
    if "last_assessment" in st.session_state:
        payload = st.session_state["last_assessment"]
        prak_pct = payload["prakriti_pct"]
        vik_pct = payload["vikriti_pct"]
        psych_pct = payload["psych_pct"]
        career = payload["career_recs"]
        rel = payload["relationship_tips"]
        health = payload["health_recs"]
        wow = payload.get("wow", {})

        st.markdown("### Results snapshot (most recent)")
        c1, c2, c3 = st.columns(3)
        c1.metric("Dominant Prakriti", max(prak_pct, key=prak_pct.get))
        c2.metric("Current Aggravation", max(vik_pct, key=vik_pct.get))
        # psych_pct max label
        try:
            max_psy = max(psych_pct, key=psych_pct.get)
        except:
            max_psy = next(iter(psych_pct.keys()))
        c3.metric("Dominant Trait", max_psy)

        # Visuals: inline radar
        radar_preview = TMP_DIR / f"preview_radar_{int(datetime.now().timestamp())}.png"
        try:
            make_radar_chart(prak_pct, vik_pct, radar_preview)
            st.image(str(radar_preview), width=360)
            try:
                radar_preview.unlink()
            except:
                pass
        except Exception:
            logger.exception("Radar preview failed")

        st.write("#### Career recommendations (top 5)")
        for r in career[:5]:
            st.write(f"- {r['role']} (score: {r['score']})  —  {r.get('reason','')}")
        st.write("#### Relationship tips")
        for t in rel:
            st.write("- " + t[0] + " — " + t[1])
        st.write("#### Health suggestions")
        st.write("Diet: " + ", ".join(health["diet"]))
        st.write("Lifestyle: " + ", ".join(health["lifestyle"]))
        st.write("Herbs: " + ", ".join(health["herbs"]))

        st.markdown("---")
        # priority strip UI
        st.write("### Priority actions")
        st.markdown(
            """
        <div style='display:flex;gap:10px;margin-bottom:12px'>
          <div style='background:#e8f7ee;padding:12px;border-radius:8px;flex:1'>
            <b>Start today</b><br>Warm water, 5 min oil rub, one focused 60–90 min block
          </div>
          <div style='background:#fff4e5;padding:12px;border-radius:8px;flex:1'>
            <b>This week</b><br>Add daily walk, add second block, start micro-project
          </div>
          <div style='background:#eef6ff;padding:12px;border-radius:8px;flex:1'>
            <b>This month</b><br>Finish & share 1 project; weekly accountability
          </div>
        </div>
        """,
            unsafe_allow_html=True,
        )

        st.write("### One-line insight (hero)")
        st.info(wow.get("hero", ""))

        if show_long_preview:
            st.write("### Life-changing 90-day plan (preview)")
            st.text(wow.get("plan", ""))
            st.write("### Habit stack (daily)")
            st.text(wow.get("habit_stack", ""))
            st.write("### One-page checklist")
            st.text(wow.get("checklist", ""))

        include_appendix = st.checkbox(
            "Include full transformation appendix in PDF", value=True
        )
        effective_wconf = WCONF.copy()
        if "pdf_wconf" in st.session_state:
            effective_wconf.update(st.session_state["pdf_wconf"])

        if st.button("Prepare Branded PDF (full report)"):
            pdf_b = branded_pdf_report(
                payload["patient"],
                prak_pct,
                vik_pct,
                psych_pct,
                career,
                rel,
                health,
                include_appendix=include_appendix,
                report_id=st.session_state.get("last_aid"),
                wconf=effective_wconf,
                wow=wow,
            )
            st.session_state["last_pdf"] = pdf_b.getvalue()
            st.success("Branded PDF prepared — download below.")
            st.balloons()
        if "last_pdf" in st.session_state:
            st.download_button(
                "Download Branded PDF (professional)",
                data=BytesIO(st.session_state["last_pdf"]),
                file_name=f"Branded_Report_{payload['patient']['name']}_{st.session_state.get('last_aid')}.pdf",
                mime="application/pdf",
            )
        else:
            # fallback directly prepare for download
            pdf_b = branded_pdf_report(
                payload["patient"],
                prak_pct,
                vik_pct,
                psych_pct,
                career,
                rel,
                health,
                include_appendix=include_appendix,
                report_id=st.session_state.get("last_aid"),
                wconf=effective_wconf,
                wow=wow,
            )
            st.download_button(
                "Download Branded PDF (professional)",
                pdf_b,
                file_name=f"Branded_Report_{payload['patient']['name']}_{st.session_state.get('last_aid')}.pdf",
                mime="application/pdf",
            )


# ---------------- DOCX generator (ensure this is at top-level, not indented) ----------------
def docx_report(
    patient,
    prakriti_pct,
    vikriti_pct,
    psych_pct,
    career_recs,
    rel_tips,
    health_recs,
    wow=None,
):
    doc = Document()
    doc.add_heading(
        f"{BRAND.get('clinic_name','Clinic')} — Personalized Report", level=1
    )
    doc.add_paragraph(
        f"Name: {patient.get('name','')}    Age: {patient.get('age','')}    Gender: {patient.get('gender','')}"
    )
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    if wow and wow.get("hero"):
        doc.add_heading("Executive one-line", level=2)
        doc.add_paragraph(wow.get("hero"))

    doc.add_heading("Prakriti (constitutional) %", level=2)
    for k, v in (prakriti_pct or {}).items():
        doc.add_paragraph(f"{k}: {v} %", style="List Bullet")

    doc.add_heading("Vikriti (today) %", level=2)
    for k, v in (vikriti_pct or {}).items():
        doc.add_paragraph(f"{k}: {v} %", style="List Bullet")

    doc.add_heading("Psychometric summary (approx)", level=2)
    for k, v in (psych_pct or {}).items():
        doc.add_paragraph(f"{k}: {v} %", style="List Bullet")

    doc.add_heading("Top career suggestions (ranked)", level=2)
    for cr in career_recs or []:
        doc.add_paragraph(
            f"{cr.get('role','Unknown')} (score: {cr.get('score','')})",
            style="List Number",
        )

    doc.add_heading("Relationship tips", level=2)
    for t in rel_tips or []:
        # tolerate tuples/lists or strings
        if isinstance(t, (list, tuple)) and len(t) >= 2:
            doc.add_paragraph(f"{t[0]} — {t[1]}", style="List Bullet")
        else:
            doc.add_paragraph(str(t), style="List Bullet")

    doc.add_heading("Health & lifestyle", level=2)
    doc.add_paragraph(
        "Diet: " + ", ".join(health_recs.get("diet", [])) if health_recs else "Diet: -"
    )
    doc.add_paragraph(
        "Lifestyle: " + ", ".join(health_recs.get("lifestyle", []))
        if health_recs
        else "Lifestyle: -"
    )
    doc.add_paragraph(
        "Herbs & cautions: " + ", ".join(health_recs.get("herbs", []))
        if health_recs
        else "Herbs & cautions: -"
    )

    if wow:
        doc.add_heading("Transformation Plan (summary)", level=2)
        doc.add_paragraph(wow.get("plan", ""))
        doc.add_heading("Daily habit stack", level=3)
        doc.add_paragraph(wow.get("habit_stack", ""))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

    docx_b = docx_report(
        payload["patient"], prak_pct, vik_pct, psych_pct, career, rel, health, wow
    )
    st.download_button(
        "Download DOCX report",
        docx_b,
        file_name=f"Report_{payload['patient']['name']}_{st.session_state.get('last_aid')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    action_pdf = onepage_actionplan_pdf(
        payload["patient"], wow.get("checklist", ""), wow.get("hero", "")
    )
    st.download_button(
        "Download 1-page Action Plan (PDF)",
        action_pdf,
        file_name=f"ActionPlan_{payload['patient']['name']}.pdf",
        mime="application/pdf",
    )

    # follow-up ICS
    ics = make_ics_followup(payload["patient"]["name"], days=7)
    st.download_button(
        "Download follow-up (.ics)",
        data=ics,
        file_name=f"followup_{payload['patient']['name']}.ics",
        mime="text/calendar",
    )


# ===== Part 3 of full app =====
# Continues the big merged file — admin + dashboard + config + utilities

# ----- Tab 3: Clinician Dashboard -----
with tabs[2]:
    st.header("Clinician Dashboard")
    st.subheader("Recent assessments")
    asses = load_assessments()
    if asses.empty:
        st.info("No assessments yet")
    else:
        st.dataframe(asses[["id", "patient_id", "assessor", "created_at"]].head(80))
        sel = st.number_input("Open assessment id", min_value=0, value=0, step=1)
        if sel > 0:
            cur.execute("SELECT data_json FROM assessments WHERE id=?", (int(sel),))
            r = cur.fetchone()
            if r:
                try:
                    st.json(json.loads(r[0]))
                except Exception:
                    st.text(r[0])
            else:
                st.warning("Not found")

    st.markdown("---")
    st.subheader("User & Admin management")
    role = st.session_state.user_info.get("role", "clinician")
    if role != "admin":
        st.info("User management features visible to admin only.")
    else:
        with st.form("create_user_form"):
            un = st.text_input("Username")
            dn = st.text_input("Display name")
            pw = st.text_input("Password", type="password")
            rrole = st.selectbox("Role", ["clinician", "admin"])
            if st.form_submit_button("Create user"):
                if not un or not pw:
                    st.warning("Provide username & password")
                else:
                    ph = pwd_context.hash(pw)
                    try:
                        cur.execute(
                            "INSERT INTO users (username, display_name, password_hash, role, created_at) VALUES (?,?,?,?,?)",
                            (un, dn, ph, rrole, datetime.now().isoformat()),
                        )
                        conn.commit()
                        st.success("User created")
                    except Exception as e:
                        st.error("Error: " + str(e))

    st.markdown("---")
    st.subheader("DB & exports")
    if st.button("Download SQLite DB"):
        data = export_sqlite_db_bytes()
        st.download_button(
            "Download DB file",
            data=data,
            file_name="ayurprakriti.db",
            mime="application/octet-stream",
        )
    st.write("Quick actions:")
    if st.button("Clear tmp files"):
        for p in TMP_DIR.glob("*"):
            try:
                p.unlink()
            except:
                pass
        st.success("Temp files removed")

# ----- Tab 4: Config & Export -----
with tabs[3]:
    st.header("Config & Export")
    st.subheader("Branding & files")
    st.write("Upload logo (PNG/JPG) and optional signature to appear on PDF.")
    logo_file = st.file_uploader("Upload logo (png/jpg)", type=["png", "jpg", "jpeg"])
    if logo_file is not None:
        save_path = APP_DIR / "logo.png"
        with open(save_path, "wb") as f:
            f.write(logo_file.getbuffer())
        st.success("Logo uploaded")
    sig_file = st.file_uploader(
        "Upload footer signature (PNG/JPG)", type=["png", "jpg", "jpeg"]
    )
    if sig_file is not None:
        sig_save = APP_DIR / "signature.png"
        with open(sig_save, "wb") as f:
            f.write(sig_file.getbuffer())
        WCONF["footer_signature_file"] = str(sig_save)
        st.success("Signature uploaded")

    st.subheader("Clinic details (branding)")
    with st.form("branding_form"):
        BRAND["clinic_name"] = st.text_input(
            "Clinic / Brand Name", value=BRAND["clinic_name"]
        )
        BRAND["tagline"] = st.text_input("Tagline", value=BRAND["tagline"])
        BRAND["doctor"] = st.text_input(
            "Doctor Name & Qualifications", value=BRAND["doctor"]
        )
        BRAND["address"] = st.text_input("Clinic Address", value=BRAND["address"])
        BRAND["phone"] = st.text_input("Phone", value=BRAND["phone"])
        BRAND["email"] = st.text_input("Email", value=BRAND["email"])
        BRAND["website"] = st.text_input("Website", value=BRAND["website"])
        BRAND["accent_color"] = st.text_input(
            "Accent color (hex)", value=BRAND["accent_color"]
        )
        if st.form_submit_button("Save branding"):
            st.success("Branding updated")

    st.markdown("---")
    st.subheader("PDF watermark & footer settings")
    col_a, col_b = st.columns(2)
    with col_a:
        wm_text = st.text_input(
            "Watermark text", value=WCONF.get("watermark_text", BRAND["clinic_name"])
        )
        wm_op = st.slider(
            "Watermark opacity (0.01 - 0.2)",
            min_value=0.01,
            max_value=0.2,
            value=float(WCONF.get("watermark_opacity", 0.06)),
            step=0.01,
        )
        show_logo = st.checkbox(
            "Show footer logo", value=WCONF.get("show_footer_logo", True)
        )
    with col_b:
        use_sig = st.checkbox(
            "Use footer signature image instead of logo",
            value=WCONF.get("use_footer_signature", False),
        )
        page_fmt = st.selectbox(
            "Page number format",
            options=["Page {page}", "Page {page} of {total}"],
            index=0,
        )
    if st.button("Save PDF/Watermark settings"):
        WCONF["watermark_text"] = wm_text
        WCONF["watermark_opacity"] = float(wm_op)
        WCONF["show_footer_logo"] = bool(show_logo)
        WCONF["use_footer_signature"] = bool(use_sig)
        WCONF["page_number_format"] = page_fmt
        st.session_state["pdf_wconf"] = WCONF.copy()
        st.success("PDF watermark & footer settings saved")

    st.markdown("---")
    st.subheader("Edit config (advanced, YAML)")
    cfg_text = yaml.safe_dump(CONFIG, sort_keys=False)
    new_cfg_text = st.text_area("Edit YAML config", cfg_text, height=300)
    if st.button("Save config file"):
        try:
            newcfg = yaml.safe_load(new_cfg_text)
            save_ok, err = save_config(newcfg)
            if save_ok:
                st.success("Config saved. Restart app to apply new questions.")
            else:
                st.error("Save failed: " + err)
        except Exception as e:
            st.error("Invalid YAML: " + str(e))

    st.markdown("---")
    st.subheader("Export & housekeeping")
    conn.commit()
    with open(DB_PATH, "rb") as f:
        dbdata = f.read()
    st.download_button(
        "Download SQLite DB",
        data=dbdata,
        file_name="ayurprakriti.db",
        mime="application/octet-stream",
    )
    st.markdown("---")
    st.caption(
        "Next steps: consider Postgres migration for multi-user access and OAuth for secure logins."
    )

# Footer small
st.write("---")
st.caption(
    f"{BRAND.get('clinic_name')} — Personalized Ayurveda Reports. Use responsibly."
)
# ===== Part 4 of full app =====
# Final utilities, cleanup, and helpful README note appended to file

# ---------------- Final helper: small README as string ----------------
APP_README = f"""
AyurPrakriti Pro Mega — single-file app
Location: {APP_DIR}
Database: {DB_PATH}
Reports: {REPORTS_DIR}
Run: streamlit run AyurPrakriti_Pro_Mega.py
Default admin: username=admin password=admin123 (change immediately)
Place DejaVuSans.ttf at {FONTS_DIR} to enable nicer PDF fonts.
"""

# show README in admin tab if user is admin
try:
    if st.session_state.user_info.get("role") == "admin":
        with st.expander("App README & paths (admin)"):
            st.code(APP_README)
except Exception:
    pass

# ---------------- Graceful message after run ----------------
logger.info("AyurPrakriti Pro Mega started at %s", datetime.now().isoformat())

# ---------------- End of file ----------------
# If you want subsequent tweaks:
# - Matching hex palette to logo: upload logo; I'll auto-suggest accent color and update BRAND.
# - Add QR link generation and online storage for shareable report.
# - Add analytics: track patient adherence and progress graphs.