# AyurPrakriti_Pro_Wow_Report.py
# Single-file Streamlit app with Cover/Hero PDF page, Radar chart, Priority action strip, and "wow" advice.

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime, timedelta
import os, json, shutil, logging, traceback
import sqlite3
from pathlib import Path
from passlib.context import CryptContext
import yaml
from docx import Document

# ReportLab
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

# ---------------- App directories & branding ----------------
APP_DIR = Path.home() / '.ayurprakriti_app'
APP_DIR.mkdir(parents=True, exist_ok=True)
FONTS_DIR = APP_DIR / 'fonts'
FONTS_DIR.mkdir(parents=True, exist_ok=True)
TMP_DIR = APP_DIR / 'tmp'
TMP_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = APP_DIR / 'ayurprakriti.db'
CFG_PATH = APP_DIR / 'config_rules.yaml'
REPORTS_DIR = APP_DIR / 'reports'
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# try to copy a logo if provided in container path (useful for your environment)
if Path('/mnt/data/logo.png').exists():
    try:
        shutil.copy('/mnt/data/logo.png', APP_DIR / 'logo.png')
    except Exception:
        pass

BRAND = {
    'clinic_name': 'Kakunje Wellness',
    'tagline': 'Authentic Ayurveda | Modern Precision',
    'doctor': 'Prof. Dr. Prasanna Kakunje, MD (Ayu), (PhD)',
    'address': 'Janani Complex, Nagarakatte Road, Moodbidri, Karnataka',
    'phone': '+91-9483697676',
    'email': 'prasanna@kakunje.com',
    'website': 'https://kakunje.com',
    'accent_color': '#0F7A61'
}

WCONF = {
    'watermark_text': BRAND['clinic_name'],
    'watermark_opacity': 0.06,
    'show_footer_logo': True,
    'use_footer_signature': False,
    'page_number_format': 'Page {page}',
    'footer_signature_file': str(APP_DIR / 'signature.png')
}

# ---------------- Logging ----------------
logger = logging.getLogger('ayurprakriti_wow')
if not logger.handlers:
    fh = logging.FileHandler(APP_DIR / 'app_debug.log')
    fh.setFormatter(logging.Formatter('%(asctime)s %(levelname)s %(message)s'))
    logger.addHandler(fh)
logger.setLevel(logging.INFO)

# ---------------- Config / default questions ----------------
pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
DEFAULT_CFG = {
    'meta': {'app_name': 'AyurPrakriti Pro', 'version': '1.0'},
    'questions': {
        'prakriti': [
            {'id': 'P1', 'text': 'Natural body frame: thin/slender', 'weights': {'Vata':1.0}},
            {'id': 'P2', 'text': 'Tendency for dry, rough skin', 'weights': {'Vata':1.0}},
            {'id': 'P3', 'text': 'Tendency for warm, reddish skin', 'weights': {'Pitta':1.0}},
            {'id': 'P4', 'text': 'Heavier, solid body build', 'weights': {'Kapha':1.0}},
            {'id': 'P5', 'text': 'Sleep depth & continuity (deep = higher)', 'weights': {'Kapha':1.0}}
        ],
        'vikriti': [
            {'id':'V1','text':'Anxiety, restlessness today','weights':{'Vata':1.0}},
            {'id':'V2','text':'Anger, irritability today','weights':{'Pitta':1.0}},
            {'id':'V3','text':'Heaviness, lethargy today','weights':{'Kapha':1.0}}
        ],
        'psychometric': [
            {'id':'E1','text':'Extraverted, enthusiastic'},{'id':'E6','text':'Reserved, quiet'},
            {'id':'A1','text':'Critical'},{'id':'A6','text':'Warm'},
            {'id':'C1','text':'Dependable'},{'id':'C6','text':'Disorganized'},
            {'id':'N1','text':'Anxious'},{'id':'N6','text':'Calm'},
            {'id':'O1','text':'Open to new experiences'},{'id':'O6','text':'Conventional'}
        ]
    },
    'mappings': {
        'career_rules': {
            'Vata': ['Writer, Designer, Creative Entrepreneur, Researcher'],
            'Pitta': ['Clinician, Analyst, Manager, Engineer'],
            'Kapha': ['Teacher, Counselor, Hospitality, Agriculture, HR']
        },
        'dosha_thresholds': {'mild':55, 'moderate':70, 'severe':85}
    }
}
if not CFG_PATH.exists():
    with open(CFG_PATH,'w', encoding='utf-8') as f:
        yaml.safe_dump(DEFAULT_CFG, f, sort_keys=False)
with open(CFG_PATH,'r', encoding='utf-8') as f:
    CONFIG = yaml.safe_load(f)

# ---------------- Database ----------------
conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
cur = conn.cursor()
cur.executescript('''
CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE, display_name TEXT, password_hash TEXT, role TEXT DEFAULT 'clinician', created_at TEXT);
CREATE TABLE IF NOT EXISTS patients (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, age INTEGER, gender TEXT, contact TEXT, created_at TEXT);
CREATE TABLE IF NOT EXISTS assessments (id INTEGER PRIMARY KEY AUTOINCREMENT, patient_id INTEGER, assessor TEXT, data_json TEXT, created_at TEXT, FOREIGN KEY(patient_id) REFERENCES patients(id));
''')
conn.commit()
cur.execute('SELECT COUNT(1) FROM users')
if cur.fetchone()[0] == 0:
    ph = pwd_context.hash('admin123')
    cur.execute('INSERT INTO users (username, display_name, password_hash, role, created_at) VALUES (?,?,?,?,?)',
                ('admin', 'Administrator', ph, 'admin', datetime.now().isoformat()))
    conn.commit()

# ---------------- Helpers: scoring & recommendations ----------------
def score_dosha_from_answers(answers, question_list):
    totals = {'Vata':0.0,'Pitta':0.0,'Kapha':0.0}
    for q in question_list:
        qid = q['id']; w = q.get('weights', {}); val = answers.get(qid, 3)
        for d in totals:
            totals[d] += w.get(d, 0) * float(val)
    s = sum(totals.values())
    if s <= 0:
        return {k: round(100/3,1) for k in totals}
    return {k: round((v/s)*100,1) for k,v in totals.items()}

def psychometric_tipiscale(answers):
    try:
        ext = (answers['E1'] + (8 - answers['E6']))/2.0
        agr = (((8 - answers['A1']) + answers['A6']))/2.0
        con = (answers['C1'] + (8 - answers['C6']))/2.0
        emo = (answers['N1'] + (8 - answers['N6']))/2.0
        ope = (answers['O1'] + (8 - answers['O6']))/2.0
    except Exception:
        return {'Extraversion':50,'Agreeableness':50,'Conscientiousness':50,'Emotionality':50,'Openness':50}
    raw = {'Extraversion': ext, 'Agreeableness': agr, 'Conscientiousness': con, 'Emotionality': emo, 'Openness': ope}
    return {k: round((v-1)/6*100,1) for k,v in raw.items()}

def recommend_career(dosha_percent, psycho_pct):
    dom = max(dosha_percent, key=dosha_percent.get)
    base = CONFIG['mappings']['career_rules'].get(dom, [])
    recs = []
    for r in base:
        score = 55
        if psycho_pct.get('Openness',50) > 60: score += 10
        if psycho_pct.get('Conscientiousness',50) > 60: score += 5
        recs.append({'role':r, 'score':score, 'reason':f"Matches {dom} with personality signals."})
    return sorted(recs, key=lambda x:-x['score'])

def recommend_relationship(dosha_pct, psycho_pct):
    tips = []
    dom = max(dosha_pct, key=dosha_pct.get)
    if dom == 'Vata':
        tips.append(('Create simple predictable rituals', 'Short, daily check-ins and predictable timing reduce anxiety.'))
    if dom == 'Pitta':
        tips.append(('Cool the tone', 'Pause 30 seconds before responding in heated moments.'))
    if dom == 'Kapha':
        tips.append(('Introduce small novelty', 'Plan a new small activity each week to spark energy.'))
    if psycho_pct.get('Agreeableness',50) < 40:
        tips.append(('Practice reflective listening', 'Repeat partner’s words once before replying.'))
    return tips

def recommend_health(dosha_pct, vikriti_pct):
    dom = max(dosha_pct, key=dosha_pct.get)
    rec = {'diet':[], 'lifestyle':[], 'herbs':[]}
    if dom == 'Vata':
        rec['diet'] = ['Warm cooked meals; add healthy oils; avoid iced drinks early.']
        rec['lifestyle'] = ['Daily warm oil self-massage (5–10 min), fixed wake/sleep times, gentle walks.']
        rec['herbs'] = ['Ashwagandha (discuss with clinician).']
    if dom == 'Pitta':
        rec['diet'] = ['Cooling foods; reduce spicy and fried items.']
        rec['lifestyle'] = ['Avoid midday heat; calming breathwork.']
        rec['herbs'] = ['Amla, Guduchi (clinician consult).']
    if dom == 'Kapha':
        rec['diet'] = ['Light, dry foods; reduce dairy & sweets.']
        rec['lifestyle'] = ['Daily brisk exercise 30–45 min; dry massage occasionally.']
        rec['herbs'] = ['Trikatu, Guggulu (clinician consult).']
    return rec

# ---------------- Radar chart ----------------
def make_radar_chart(prakriti, vikriti, filename: Path, title='Prakriti vs Vikriti'):
    labels = list(prakriti.keys())
    n = len(labels)
    angles = np.linspace(0, 2*np.pi, n, endpoint=False).tolist()
    vals1 = [prakriti[l] for l in labels]
    vals2 = [vikriti.get(l, 0) for l in labels]
    vals1 += vals1[:1]; vals2 += vals2[:1]; angles += angles[:1]
    fig = plt.figure(figsize=(4.2,4.2)); ax = fig.add_subplot(111, polar=True)
    ax.set_theta_offset(np.pi/2); ax.set_theta_direction(-1)
    ax.plot(angles, vals1, linewidth=2, label='Prakriti'); ax.fill(angles, vals1, alpha=0.25)
    ax.plot(angles, vals2, linewidth=2, label='Vikriti'); ax.fill(angles, vals2, alpha=0.12)
    ax.set_thetagrids(np.degrees(angles[:-1]), labels)
    ax.set_ylim(0,100)
    ax.set_title(title, pad=10)
    ax.legend(loc='upper right', bbox_to_anchor=(1.3,1.1))
    plt.tight_layout(); fig.savefig(filename, dpi=150); plt.close(fig)

# ---------------- Fonts (DejaVu) registration if available ----------------
DEJAVU_PATH = None
_fonts = list(FONTS_DIR.glob("DejaVuSans*.ttf"))
if _fonts:
    DEJAVU_PATH = str(_fonts[0])
else:
    for cand in [r"C:\Windows\Fonts\DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/Library/Fonts/DejaVuSans.ttf"]:
        if os.path.exists(cand):
            DEJAVU_PATH = cand; break
if DEJAVU_PATH:
    try:
        pdfmetrics.registerFont(TTFont('DejaVuSans', DEJAVU_PATH))
    except Exception:
        logger.exception("Font register failed")

# ---------------- "Wow" plain-language / life-changing advice generator ----------------
def generate_wow_advice(patient, prakriti_pct, vikriti_pct, psych_pct, career_recs, rel_tips, health_recs):
    dom = max(prakriti_pct, key=prakriti_pct.get)
    current = max(vikriti_pct, key=vikriti_pct.get)
    # One-line charismatic insight (hero)
    hero = (f"{patient.get('name','You')} — creative energy + steady rituals = powerful results. "
            f"You're primarily {dom}. Right now you may feel {('slow & heavy' if current=='Kapha' else 'scattered & anxious' if current=='Vata' else 'hot & impatient')}.")
    # 90-day identity+behavior plan (simple steps)
    plan_lines = [
        "90-day transformation plan (small daily actions that become identity):",
        "1) Identity pledge (Day 1): Write one line: 'I am someone who finishes what they start with calm focus.' Put it on your phone wallpaper.",
        "2) Daily core ritual (0–21 days): Warm water, 5–10 min oil massage or warming shower, 2 focused work blocks (60–90 min).",
        "3) Weekly mastery (weeks 3–12): Finish one small creative project every 2–3 weeks and share it (blog, social, colleague).",
        "4) Accountability (start now): Pick one friend/peer to check progress weekly for 12 weeks (2-min message).",
        "5) Measure progress: Record morning energy (1–5) and sleep times daily for 90 days; review at day 14, 45, 90.",
        "6) Recalibrate: If sleep/energy don't improve after 14 days, book a short consult — tweak herbs/doses or routines."
    ]
    plan = "\n".join(plan_lines)
    # Life-changing single habit stack (keep very small)
    habit_stack = [
        "Life-changing habit stack (do these in order, takes 15–25 minutes total):",
        "A) Warm water on waking + 2 min breathing (inhale 4s, exhale 6s)",
        "B) 5–10 min oil self-massage or 10 min full-body stretching",
        "C) One 60–90 min focused creative/work block (timer on)",
        "D) Evening: reflect 2 things done, prepare 1 clear task for next day"
    ]
    habit_stack_text = "\n".join(habit_stack)
    # Concrete "wow" tips (actionable, emotionally resonant)
    wow_tips = [
        "- If you want more calm: reduce decision load — only 3 choices for morning clothes/breakfast.",
        "- If you want more creative output: ship one small thing per week and celebrate it.",
        "- If you want better relationships: do a 3-minute daily gratitude note for a partner or colleague.",
        "- If you want better health: commit to 21 days of the routine — habits stick around that mark."
    ]
    wow_text = "\n".join(wow_tips)
    # One-page checklist (copyable)
    checklist = [
        "ONE-PAGE ACTION CHECKLIST",
        "- Morning: warm water + 2 min breathing + 5–10 min oil rub/stretch",
        "- Work: two focused work blocks (60–90 min each). Timer ON.",
        "- Movement: daily 25–35 min walk or light yoga.",
        "- Evening: light dinner by 8 pm, reflect on 2 wins.",
        "- Weekly: publish/share 1 small creative item; 20-min planning on Sunday.",
        "- Accountability: weekly check-in with your chosen peer for 12 weeks."
    ]
    checklist_text = "\n".join(checklist)
    # Doctor's short signed note (friendly)
    doctor_note = ("Doctor's note: Start the 'Start Today' items now. "
                   "Small, consistent steps matter more than big, rare interventions. "
                   "We'll review your progress at 2 weeks and tune the plan.")
    return {'hero': hero, 'plan': plan, 'habit_stack': habit_stack_text, 'wow_text': wow_text, 'checklist': checklist_text, 'doctor_note': doctor_note}

# ---------------- One-page action plan PDF ----------------
def onepage_actionplan_pdf(patient, checklist_text, hero_text):
    buf = BytesIO(); c = canvas.Canvas(buf, pagesize=A4)
    left = 20*mm; y = A4[1] - 30*mm
    try:
        if DEJAVU_PATH:
            c.setFont('DejaVuSans', 14)
        else:
            c.setFont('Helvetica-Bold', 14)
    except Exception:
        c.setFont('Helvetica-Bold', 14)
    c.drawString(left, y, BRAND['clinic_name']); y -= 9*mm
    c.setFont('Helvetica', 10); c.drawString(left, y, hero_text); y -= 10*mm
    c.setFont('Helvetica', 10)
    for line in checklist_text.split('\n'):
        if not line.strip(): continue
        if line.startswith('- '):
            c.drawString(left + 4*mm, y, u'\u2022 ' + line[2:])
        else:
            c.drawString(left, y, line)
        y -= 7*mm
        if y < 30*mm:
            c.showPage(); y = A4[1] - 30*mm
    c.setFont('Helvetica', 8); c.drawString(left, 12*mm, f"{BRAND['clinic_name']} — {BRAND['phone']} — {BRAND['email']}")
    c.save(); buf.seek(0); return buf

# ---------------- Simple fallback wrapper for text ----------------
def _wrap_text_simple(text, chars_per_line=95):
    words = str(text).split(); lines=[]; cur=""
    for w in words:
        if len(cur)+len(w)+1 <= chars_per_line: cur = cur + (" " if cur else "") + w
        else: lines.append(cur); cur=w
    if cur: lines.append(cur)
    return lines

# ---------------- Branded PDF with Cover, Radar, Priority strip ----------------
def branded_pdf_report(patient, prakriti_pct, vikriti_pct, psych_pct, career_recs, rel_tips, health_recs,
                       include_appendix=False, report_id=None, wconf=None, wow=None):
    if wconf is None: wconf = WCONF
    try:
        radar_path = TMP_DIR / f"radar_{int(datetime.now().timestamp())}.png"
        make_radar_chart(prakriti_pct, vikriti_pct, radar_path, title='Prakriti vs Vikriti')
    except Exception:
        radar_path = None
    try:
        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm)
        styles = getSampleStyleSheet()
        base_font = 'DejaVuSans' if DEJAVU_PATH else 'Helvetica'
        accent = colors.HexColor(BRAND.get('accent_color','#0F7A61'))
        styles.add(ParagraphStyle(name='Hero', fontName=base_font, fontSize=20, leading=24, alignment=1, spaceAfter=6))
        styles.add(ParagraphStyle(name='HeroSub', fontName=base_font, fontSize=11, leading=13, alignment=1, textColor=colors.darkgrey))
        styles.add(ParagraphStyle(name='Section', fontName=base_font, fontSize=12, leading=14, textColor=accent))
        styles.add(ParagraphStyle(name='Body', fontName=base_font, fontSize=10, leading=12))
        flow = []
        # Cover / Hero page
        flow.append(Spacer(1,6))
        # big name
        flow.append(Paragraph(patient.get('name','Patient Name'), styles['Hero']))
        if wow and wow.get('hero'):
            flow.append(Paragraph(wow['hero'], styles['HeroSub']))
        flow.append(Spacer(1,6))
        # badges row (Dominant / Current / Top career)
        badges = [
            Paragraph(f"<b>Dominant</b><br/>{max(prakriti_pct,key=prakriti_pct.get)}", styles['Body']),
            Paragraph(f"<b>Current</b><br/>{max(vikriti_pct,key=vikriti_pct.get)}", styles['Body']),
            Paragraph(f"<b>Top career</b><br/>{career_recs[0]['role'] if career_recs else '-'}", styles['Body'])
        ]
        table = Table([[badges[0], badges[1], badges[2]]], colWidths=[60*mm,60*mm,60*mm])
        table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.whitesmoke),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('ALIGN',(0,0),(-1,-1),'CENTER')]))
        flow.append(table)
        flow.append(Spacer(1,8))
        # radar chart on cover if available
        if radar_path and radar_path.exists():
            try:
                img = RLImage(str(radar_path), width=120*mm, height=120*mm)
                flow.append(img)
            except Exception:
                pass
        flow.append(Spacer(1,8))
        # Doctor note + signature (if available)
        doc_note = wow.get('doctor_note') if wow else ''
        if doc_note:
            flow.append(Paragraph(f"<i>{doc_note}</i>", styles['Body']))
            sig_path = APP_DIR / 'signature.png'
            if sig_path.exists():
                try:
                    sig_img = scaled_rl_image(sig_path, max_w_mm=40, max_h_mm=20)
                    flow.append(sig_img)
                except Exception:
                    pass
        flow.append(PageBreak())

        # Main content - summary, charts, tables
        flow.append(Paragraph("Executive summary", styles['Section']))
        flow.append(Paragraph("Below are clear findings and prioritized actions to start today.", styles['Body']))
        flow.append(Spacer(1,6))
        # charts area (radar + bar charts)
        if radar_path and radar_path.exists():
            try:
                img = RLImage(str(radar_path), width=120*mm, height=120*mm)
                flow.append(img); flow.append(Spacer(1,6))
            except Exception:
                pass

        # Priority action strip (Start today / This week / This month)
        priority = [
            ("Start today", "Warm water on waking; 5–10 min warm oil rub or stretch; one focused 60–90 min block"),
            ("This week", "Add a second focused block; daily 20–35 min walk; start a small micro-project"),
            ("This month", "Finish + share one creative project; set a weekly accountability check-in")
        ]
        cols_row = []
        for title, text in priority:
            p = Paragraph(f"<b>{title}</b><br/>{text}", styles['Body'])
            cols_row.append(p)
        strip = Table([cols_row], colWidths=[60*mm,60*mm,60*mm])
        strip.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,-1),colors.Color(0.96,0.98,0.95)),
            ('BOX',(0,0),(-1,-1),0.5,colors.lightgrey),
            ('VALIGN',(0,0),(-1,-1),'TOP'),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ('LEFTPADDING',(0,0),(-1,-1),6), ('RIGHTPADDING',(0,0),(-1,-1),6)
        ]))
        flow.append(strip); flow.append(Spacer(1,8))

        # Recommendations: Career / Relationship / Health short blocks
        flow.append(Paragraph("Top recommendations (short)", styles['Section']))
        flow.append(Paragraph("Career: " + (career_recs[0]['role'] if career_recs else '—'), styles['Body']))
        flow.append(Paragraph("Relationship: " + (rel_tips[0][1] if rel_tips else '—'), styles['Body']))
        flow.append(Paragraph("Health: " + (', '.join(health_recs.get('diet',[])) if health_recs else '—'), styles['Body']))
        flow.append(Spacer(1,8))

        # Include long plain-language/wow plan as appendix if requested
        if include_appendix and wow:
            flow.append(PageBreak())
            flow.append(Paragraph("APPENDIX — Transformation Plan", styles['Section']))
            flow.append(Spacer(1,6))
            flow.append(Paragraph("<b>90-day transformation (practical)</b>", styles['Body']))
            for line in wow['plan'].split('\n'):
                flow.append(Paragraph(line, styles['Body']))
            flow.append(Spacer(1,6))
            flow.append(Paragraph("<b>Habit stack (daily)</b>", styles['Body']))
            for line in wow['habit_stack'].split('\n'):
                flow.append(Paragraph(line, styles['Body']))
            flow.append(Spacer(1,6))
            flow.append(Paragraph("<b>Concrete wow tips</b>", styles['Body']))
            for line in wow['wow_text'].split('\n'):
                flow.append(Paragraph(line, styles['Body']))
            flow.append(Spacer(1,6))
            flow.append(Paragraph("<b>One-page checklist</b>", styles['Body']))
            for line in wow['checklist'].split('\n'):
                flow.append(Paragraph(line, styles['Body']))

        # footer contact block
        flow.append(PageBreak())
        flow.append(Paragraph(f"{BRAND['clinic_name']} — {BRAND['doctor']} — {BRAND['phone']} — {BRAND['email']}", styles['Body']))

        # page watermark & footer
        def _draw_footer_and_watermark(canvas_obj, doc_obj):
            try:
                canvas_obj.saveState()
                W,H = A4
                try:
                    canvas_obj.setFont('DejaVuSans', 40) if DEJAVU_PATH else canvas_obj.setFont('Helvetica-Bold', 40)
                except Exception:
                    canvas_obj.setFont('Helvetica-Bold', 40)
                opacity = float(wconf.get('watermark_opacity', 0.06))
                try:
                    canvas_obj.setFillAlpha(opacity)
                except Exception:
                    canvas_obj.setFillColorRGB(0.7,0.7,0.7)
                canvas_obj.translate(W/2.0, H/2.0); canvas_obj.rotate(30)
                canvas_obj.drawCentredString(0, 0, wconf.get('watermark_text', BRAND['clinic_name']))
                canvas_obj.restoreState()
            except Exception:
                logger.exception("watermark failed")
            try:
                canvas_obj.saveState()
                footer_y = 15*mm
                canvas_obj.setStrokeColor(colors.lightgrey); canvas_obj.setLineWidth(0.5)
                canvas_obj.line(18*mm, footer_y + 8, (A4[0]-18*mm), footer_y + 8)
                logo_path_local = APP_DIR / 'logo.png'
                x = 20*mm
                if wconf.get('show_footer_logo', True) and logo_path_local.exists():
                    try:
                        reader = ImageReader(str(logo_path_local)); iw, ih = reader.getSize(); target_h = 10*mm; scale = target_h/ih
                        canvas_obj.drawImage(str(logo_path_local), x, footer_y - 2, width=iw*scale, height=ih*scale, mask='auto'); x += (iw*scale) + 4
                    except Exception:
                        pass
                canvas_obj.setFont('Helvetica', 8)
                canvas_obj.drawString(x, footer_y, f"{BRAND['clinic_name']} — {BRAND['phone']} — {BRAND['email']}")
                page_num = canvas_obj.getPageNumber()
                canvas_obj.drawRightString(A4[0] - 18*mm, footer_y, wconf.get('page_number_format','Page {page}').format(page=page_num))
                canvas_obj.restoreState()
            except Exception:
                logger.exception("footer failed")
        doc.build(flow, onFirstPage=_draw_footer_and_watermark, onLaterPages=_draw_footer_and_watermark)
        buf.seek(0)
        if radar_path and radar_path.exists():
            try: radar_path.unlink()
            except: pass
        return buf
    except Exception as e:
        logger.exception("Branded PDF failed: %s", e)
        # fallback simple canvas
        buf = BytesIO(); c = canvas.Canvas(buf, pagesize=A4)
        c.drawString(50,800, f"{BRAND['clinic_name']} - Simple fallback report")
        c.save(); buf.seek(0); return buf

# ---------------- Docx ----------------
def docx_report(patient, prakriti_pct, vikriti_pct, psych_pct, career_recs, rel_tips, health_recs, wow):
    doc = Document()
    doc.add_heading(f"{BRAND['clinic_name']} — Personalized Report", level=1)
    doc.add_paragraph(f"Name: {patient.get('name')}    Age: {patient.get('age')}    Gender: {patient.get('gender')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading('Executive one-line', level=2)
    doc.add_paragraph(wow.get('hero',''))
    doc.add_heading('90-day plan', level=2)
    for line in wow.get('plan','').split('\n'): doc.add_paragraph(line)
    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio

# ---------------- Streamlit UI ----------------
st.set_page_config(page_title="AyurPrakriti Pro — Wow Reports", layout='wide')
st.markdown("<style>section[data-testid='stSidebar'] {background-color: #f7f7fa}</style>", unsafe_allow_html=True)
col1, col2 = st.columns([6,4])
with col1:
    st.title("AyurPrakriti Pro — Wow Reports")
    st.caption("Professional, beautiful personalized Ayurveda reports")
with col2:
    st.write(datetime.now().strftime('%Y-%m-%d %H:%M'))

# Sidebar auth
st.sidebar.subheader('Login')
username = st.sidebar.text_input('Username'); password = st.sidebar.text_input('Password', type='password')
if 'auth' not in st.session_state: st.session_state.auth = False
if st.sidebar.button('Login'):
    cur.execute('SELECT password_hash, display_name, role FROM users WHERE username=?', (username,))
    r = cur.fetchone()
    if r and pwd_context.verify(password, r[0]):
        st.session_state.auth = True; st.session_state.user = username; st.session_state.user_info = {'display_name': r[1], 'role': r[2]}
        st.sidebar.success(f"Welcome {r[1]}")
    else:
        st.sidebar.error('Invalid credentials')
if not st.session_state.auth:
    st.info('Please login in the sidebar (default admin / admin123)'); st.stop()

tabs = st.tabs(['Patients','New Assessment','Dashboard','Config & Export'])

# Patients
with tabs[0]:
    st.header('Patients')
    with st.expander('Create new patient'):
        with st.form('pform'):
            name = st.text_input('Full name'); age = st.number_input('Age', min_value=0, max_value=120, value=30)
            gender = st.selectbox('Gender',['Male','Female','Other','Prefer not to say']); contact = st.text_input('Contact')
            if st.form_submit_button('Create'):
                if not name: st.warning('Name required')
                else:
                    cur.execute('INSERT INTO patients (name, age, gender, contact, created_at) VALUES (?,?,?,?,?)',
                                (name, age, gender, contact, datetime.now().isoformat())); conn.commit()
                    st.success('Patient created')
    patients_df = pd.read_sql_query('SELECT * FROM patients ORDER BY created_at DESC', conn)
    st.dataframe(patients_df)

# New Assessment
with tabs[1]:
    st.header('New Assessment — Prakriti / Vikriti / Psychometrics')
    patients = pd.read_sql_query('SELECT * FROM patients ORDER BY created_at DESC', conn)
    if patients.empty:
        st.info('Create a patient first'); st.stop()
    psel = st.selectbox('Select patient', options=patients['id'].tolist(), format_func=lambda x: f"{int(x)} - {patients.loc[patients['id']==x,'name'].values[0]}")
    patient_row = patients[patients['id']==psel].iloc[0].to_dict()
    st.markdown(f"**Patient:** {patient_row['name']} — Age {patient_row['age']} — {patient_row['gender']}")
    st.markdown('---')
    # Questions
    pr_qs = CONFIG['questions']['prakriti']; pr_answers={}
    cols = st.columns(2)
    for i,q in enumerate(pr_qs):
        with cols[i%2]: pr_answers[q['id']] = st.slider(q['text'], 1, 5, 3, key=f"pr_{q['id']}")
    vk_qs = CONFIG['questions']['vikriti']; vk_answers={}
    cols = st.columns(3)
    for i,q in enumerate(vk_qs):
        with cols[i%3]: vk_answers[q['id']] = st.slider(q['text'],1,5,1,key=f"vk_{q['id']}")
    psy_qs = CONFIG['questions']['psychometric']; psy_answers={}
    cols = st.columns(2)
    for i,q in enumerate(psy_qs):
        with cols[i%2]: psy_answers[q['id']] = st.slider(q['text'],1,7,4,key=f"psy_{q['id']}")
    if st.button('Compute & Save'):
        prak_pct = score_dosha_from_answers(pr_answers, pr_qs)
        vik_pct = score_dosha_from_answers(vk_answers, vk_qs)
        psych_pct = psychometric_tipiscale(psy_answers)
        career = recommend_career(prak_pct, psych_pct)
        rel = recommend_relationship(prak_pct, psych_pct)
        health = recommend_health(prak_pct, vik_pct)
        payload = {'patient': patient_row, 'prakriti_pct': prak_pct, 'vikriti_pct': vik_pct, 'psych_pct': psych_pct,
                   'career_recs': career, 'relationship_tips': rel, 'health_recs': health, 'created_at': datetime.now().isoformat()}
        aid = cur.execute('INSERT INTO assessments (patient_id, assessor, data_json, created_at) VALUES (?,?,?,?)',
                          (patient_row['id'], st.session_state.user, json.dumps(payload, ensure_ascii=False), datetime.now().isoformat()))
        conn.commit()
        # wow advice
        wow = generate_wow_advice(patient_row, prak_pct, vik_pct, psych_pct, career, rel, health)
        payload['wow'] = wow
        st.session_state['last_assessment'] = payload
        st.session_state['last_aid'] = cur.lastrowid
        st.success('Assessment saved')

    if 'last_assessment' in st.session_state:
        payload = st.session_state['last_assessment']
        prak_pct = payload['prakriti_pct']; vik_pct = payload['vikriti_pct']; psych_pct = payload['psych_pct']
        career = payload['career_recs']; rel = payload['relationship_tips']; health = payload['health_recs']; wow = payload.get('wow', {})
        # badges UI
        b1,b2,b3 = st.columns(3)
        b1.markdown(f"<div style='background:#e8f7ee;padding:12px;border-radius:8px'><h3 style='margin:0'>{max(prak_pct,key=prak_pct.get)}</h3><small>Dominant</small></div>", unsafe_allow_html=True)
        b2.markdown(f"<div style='background:#fff4e5;padding:12px;border-radius:8px'><h3 style='margin:0'>{max(vik_pct,key=vik_pct.get)}</h3><small>Current</small></div>", unsafe_allow_html=True)
        b3.markdown(f"<div style='background:#eef6ff;padding:12px;border-radius:8px'><h3 style='margin:0'>{career[0]['role'] if career else '-'}</h3><small>Top career match</small></div>", unsafe_allow_html=True)
        st.markdown('---')
        st.write('### Quick snapshot') 
        c1,c2,c3 = st.columns(3)
        c1.metric('Dominant', max(prak_pct, key=prak_pct.get))
        c2.metric('Current', max(vik_pct, key=vik_pct.get))
        c3.metric('Top career', career[0]['role'] if career else '-')
        st.write('### Visuals')
        # show radar chart inline
        radar_local = TMP_DIR / 'preview_radar.png'
        try:
            make_radar_chart(prak_pct, vik_pct, radar_local)
            st.image(str(radar_local), use_column_width=False, width=360)
            try: radar_local.unlink()
            except: pass
        except Exception:
            pass
        st.write('### Start today (priority)')
        st.markdown("""
        <div style='display:flex;gap:10px'>
          <div style='background:#e8f7ee;padding:12px;border-radius:8px;flex:1'>
            <b>Start today</b><br>Warm water, 5 min oil rub, one focused 60–90 min block
          </div>
          <div style='background:#fff4e5;padding:12px;border-radius:8px;flex:1'>
            <b>This week</b><br>Add daily walk, add second block, start micro-project
          </div>
          <div style='background:#eef6ff;padding:12px;border-radius:8px;flex:1'>
            <b>This month</b><br>Finish & share 1 project; weekly accountability
          </div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('---')
        st.write('### One-line insight (hero)')
        st.info(wow.get('hero',''))
        st.write('### Life-changing plan (summary)')
        st.write(wow.get('plan',''))
        st.write('### Habit stack (daily)')
        st.write(wow.get('habit_stack',''))
        st.markdown('---')

        include_appendix = st.checkbox('Include full transformation appendix in PDF', value=True)
        if st.button('Prepare Branded PDF (wow report)'):
            pdf_b = branded_pdf_report(payload['patient'], prak_pct, vik_pct, psych_pct, career, rel, health,
                                      include_appendix=include_appendix, report_id=st.session_state.get('last_aid'), wconf=WCONF, wow=wow)
            st.session_state['last_pdf'] = pdf_b.getvalue()
            st.success('PDF prepared — download below'); st.balloons()
        if 'last_pdf' in st.session_state:
            st.download_button('Download wow Branded PDF', data=BytesIO(st.session_state['last_pdf']),
                               file_name=f"WowReport_{patient_row['name']}_{st.session_state.get('last_aid',0)}.pdf", mime='application/pdf')
        # 1-page action plan
        action_pdf = onepage_actionplan_pdf(payload['patient'], wow.get('checklist',''), wow.get('hero',''))
        st.download_button('Download 1-page Action Plan (PDF)', action_pdf, file_name=f"ActionPlan_{patient_row['name']}.pdf", mime='application/pdf')
        # docx
        docx_b = docx_report(payload['patient'], prak_pct, vik_pct, psych_pct, career, rel, health, wow)
        st.download_button('Download DOCX report', docx_b, file_name=f"Report_{patient_row['name']}.docx",
                           mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Dashboard
with tabs[2]:
    st.header('Clinician Dashboard')
    df = pd.read_sql_query('SELECT * FROM assessments ORDER BY created_at DESC', conn)
    if df.empty: st.info('No assessments yet')
    else:
        st.dataframe(df[['id','patient_id','assessor','created_at']].head(40))
        sel = st.number_input('Open assessment id', min_value=0, value=0, step=1)
        if sel > 0:
            cur.execute('SELECT data_json FROM assessments WHERE id=?', (int(sel),))
            r = cur.fetchone()
            if r: st.json(json.loads(r[0])); else: st.warning('Not found')
    st.markdown('---')
    if st.session_state.user_info.get('role') == 'admin':
        with st.form('create_user'):
            un = st.text_input('Username'); dn = st.text_input('Display name'); pw = st.text_input('Password', type='password')
            role = st.selectbox('Role',['clinician','admin'])
            if st.form_submit_button('Create user'):
                if not un or not pw: st.warning('Provide username & password')
                else:
                    ph = pwd_context.hash(pw)
                    try:
                        cur.execute('INSERT INTO users (username, display_name, password_hash, role, created_at) VALUES (?,?,?,?,?)',
                                    (un, dn, ph, role, datetime.now().isoformat())); conn.commit(); st.success('User created')
                    except Exception as e: st.error(str(e))

# Config & Export
with tabs[3]:
    st.header('Config & Export')
    st.subheader('Upload logo (appears in footer & cover)')
    logo_file = st.file_uploader('Logo PNG/JPG', type=['png','jpg','jpeg'])
    if logo_file:
        s = APP_DIR / 'logo.png'
        with open(s,'wb') as f: f.write(logo_file.getbuffer())
        st.success('Logo saved')
    st.subheader('PDF settings')
    wm = st.text_input('Watermark text', value=WCONF.get('watermark_text',BRAND['clinic_name']))
    wm_op = st.slider('Watermark opacity', 0.01, 0.2, float(WCONF.get('watermark_opacity',0.06)))
    show_logo = st.checkbox('Show footer logo', value=WCONF.get('show_footer_logo', True))
    if st.button('Save PDF settings'):
        WCONF['watermark_text'] = wm; WCONF['watermark_opacity'] = float(wm_op); WCONF['show_footer_logo'] = bool(show_logo)
        st.session_state['pdf_wconf'] = WCONF.copy(); st.success('Settings saved')

st.write('---')
st.caption('Designed for Kakunje Wellness — personalized, actionable Ayurveda reports.')

# End
